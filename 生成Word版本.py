#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将SRT字幕文件转换为Word文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def srt_to_docx(srt_file, docx_file):
    """
    将SRT字幕文件转换为Word文档
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('流媒体与流行明星 - 字幕稿', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 读取SRT文件
    with open(srt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 解析SRT内容
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 检查是否是序号
        if line.isdigit():
            subtitle_num = line
            # 读取时间戳
            i += 1
            timestamp = lines[i].strip()
            # 读取内容
            i += 1
            content = lines[i].strip()
            # 读取可能的多行内容
            while i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].strip().isdigit():
                i += 1
                content += '\n' + lines[i].strip()
            
            # 添加到文档
            # 添加时间戳（灰色、较小字体）
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(timestamp)
            timestamp_run.font.size = Pt(9)
            timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
            timestamp_para.paragraph_format.space_before = Pt(6)
            timestamp_para.paragraph_format.space_after = Pt(2)
            
            # 添加内容（正常大小）
            content_para = doc.add_paragraph(content)
            content_para.paragraph_format.space_after = Pt(12)
            
            i += 1
        else:
            i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"✓ Word文档已生成: {docx_file}")

if __name__ == '__main__':
    srt_file = '/Users/ilovemeiya/🇫🇷法语翻译/ilovemeiya/字幕_中文.srt'
    docx_file = '/Users/ilovemeiya/🇫🇷法语翻译/ilovemeiya/字幕_中文.docx'
    srt_to_docx(srt_file, docx_file)
